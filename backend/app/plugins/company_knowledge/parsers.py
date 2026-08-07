"""电子版公司资料的文档解析适配器（PDF / Word → Markdown）。

对应迭代文档"第二步：文档解析"，用本地工具解析，不消耗 LLM token：
- PDF：PyMuPDF（fitz），逐页提取文本
- Word .docx：python-docx，段落 + 表格 → Markdown
解析结果进入与 TXT/Markdown 相同的后续流程（Markdown 审核 → 数据预处理 → 切分）。
"""

from dataclasses import dataclass
from pathlib import Path


class DocumentParseError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    source_format: str
    file_name: str
    content: str
    warnings: list[str]


def parse_document(file_name: str, content: bytes) -> ParsedDocument:
    """按文件后缀分发到对应解析器，返回 Markdown 正文与转换告警。"""
    suffix = Path(file_name).suffix.lower().lstrip(".")
    if suffix == "pdf":
        return _parse_pdf(file_name, content)
    if suffix == "docx":
        return _parse_docx(file_name, content)
    raise DocumentParseError(f"不支持的文档格式：.{suffix}")


def _parse_pdf(file_name: str, content: bytes) -> ParsedDocument:
    try:
        import pymupdf  # PyMuPDF 1.24+ 推荐入口；旧版本可用 fitz
    except ImportError:
        import fitz as pymupdf  # type: ignore[no-redef]

    try:
        document = pymupdf.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise DocumentParseError("PDF 解析失败，请确认文件未损坏") from exc

    try:
        pages: list[str] = []
        for page in document:
            text = page.get_text("text").strip()
            if text:
                pages.append(text)
    except Exception as exc:
        raise DocumentParseError("PDF 文本提取失败") from exc
    finally:
        document.close()

    if not pages:
        raise DocumentParseError("PDF 未提取到可索引文本，可能是扫描件，暂不支持 OCR")

    warnings: list[str] = []
    if len(pages) == 1:
        warnings.append("PDF 仅含 1 页，请确认内容完整。")
    if any(_looks_like_scanned(page) for page in pages):
        warnings.append("部分页面文本极少，可能是扫描图片页，内容可能不完整。")

    content = "\n\n".join(pages)
    return ParsedDocument(source_format="pdf", file_name=file_name, content=content, warnings=warnings)


def _parse_docx(file_name: str, content: bytes) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentParseError("Word 解析组件未安装") from exc

    import io

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise DocumentParseError("Word 解析失败，请确认文件为 .docx 格式且未损坏") from exc

    blocks: list[str] = []
    warnings: list[str] = []

    # 遍历文档 body 子元素，按阅读顺序输出段落与表格
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    table_warned = False
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            blocks.append(_table_to_markdown(table))
            if not table_warned:
                warnings.append("Word 表格已转换为 Markdown 表格，请确认列对齐。")
                table_warned = True

    if not blocks:
        raise DocumentParseError("Word 未提取到可索引文本")

    content = "\n\n".join(blocks)
    return ParsedDocument(source_format="docx", file_name=file_name, content=content, warnings=warnings)


def _table_to_markdown(table) -> str:
    rows: list[str] = []
    for row_index, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if row_index == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def _looks_like_scanned(page_text: str) -> bool:
    """启发式：一页文本少于 50 字符视为疑似扫描图片页。"""
    return len(page_text.strip()) < 50
